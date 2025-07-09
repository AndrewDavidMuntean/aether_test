import streamlit as st
import urllib.parse
import gdown
from PIL import Image
import torch
import requests
from transformers import CLIPProcessor, CLIPModel
from supabase import create_client, Client
from docx import Document
from io import BytesIO
import os
from dotenv import load_dotenv
from docx import Document
from transformers import AutoTokenizer, AutoModel

from transformers import AutoModelForCausalLM




import torch
import torch.nn as nn
from torchvision.models import densenet121
from torchvision import transforms
from PIL import Image

#Model definition

class VisionLanguageModel(nn.Module):
    def __init__(self):
        super().__init__()
        # Vision encoder
        self.encoder = densenet121(weights=None)
        self.encoder.classifier = nn.Identity()

        # Text embedding projection (placeholder)
        self.text_proj = nn.Linear(256, 1024)

        # Final fusion + decision layer
        self.fc = nn.Linear(1024, 1)

    def forward(self, image_tensor, text_embedding):
        vision_feat = self.encoder(image_tensor)  # [batch, 1024]
        text_feat = self.text_proj(text_embedding)  # [batch, 1024]
        fused = vision_feat + text_feat  # Simple fusion; replace with better logic if needed
        out = self.fc(fused)
        return out


class_names = [
    "copd signs",
    "cardiomegaly",
    "aortic elongation",
    "pleural effusion",
    "pneumonia",
    "interstitial pattern",
    "air trapping",
    "alveolar pattern",
    "laminar atelectasis",
    "infiltrates",
    "apical pleural thickening"
]


MODEL_ID = "1ZfakZ-a7GLC7WJDwuqtRviitZpKxAj9E"
MODEL_URL = f"https://drive.google.com/uc?id={MODEL_ID}"
MODEL_PATH = "fold1_best_model_24_epochs.pt"

# Download model if not present
if not os.path.exists(MODEL_PATH):
    print("Downloading model with gdown...")
    gdown.download(MODEL_URL, MODEL_PATH, quiet=False)
    print("Download complete.")
else:
    print("Model already exists.")


# Load model
checkpoint = torch.load(MODEL_PATH, map_location="cpu")
print("Checkpoint keys:", checkpoint.keys())
print("Model state_dict keys (sample):", list(checkpoint["model_state_dict"].keys())[:10])

model = VisionLanguageModel()
model.load_state_dict(checkpoint["model_state_dict"], strict=False)
model.eval()
print("Model loaded successfully!")




tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
text_encoder = AutoModel.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")





load_dotenv()
SUPABASE_URL = "https://bhewsjowxlgxxzeorwbj.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImJoZXdzam93eGxneHh6ZW9yd2JqIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDg1MTIxNTcsImV4cCI6MjA2NDA4ODE1N30.b7Ccc-v4xqNKsjyNxcna1i4ePjRDTTUSWVraJxMCttU"


supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)



for key in ["user", "access_token", "pastquestions", "past_responses", "selected_question", "ai_response"]:
    if key not in st.session_state:
        st.session_state[key] = None if key in ["user", "access_token"] else []

def login():
    email = st.text_input("Email")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        try:
            result = supabase.auth.sign_in_with_password({"email": email, "password": password})
            st.session_state.user = result.user
            st.session_state.access_token = result.session.access_token
            supabase.postgrest.auth(st.session_state.access_token)
            st.success("Logged in successfully!")
        except Exception as e:
            st.error(f"Login failed: {e}")


def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return full_text












def doctors_report_summary(docx_file):
    api_key = "C7Xwjta0ERFDWptomYooHrTOqMDWlS51"
    report_text = extract_text_from_docx(docx_file)

    question = "Summarize this radiology report in three bullet points."
    prompt = f"The following is a detailed radiology report:\n\n{report_text}\n\n{question}"

    
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"model": "mistral-small", "messages": [{"role": "user", "content": prompt}]}
    response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]



def get_model_response(image: Image.Image, question: str, model: nn.Module):
    

    labels_above_threshold = [label for label, score in st.session_state.predictions if score >= 0.7] #THRESHOLD DETERMINED HERE!!!!!!!!!!!!!!!!!
    label_text = ", ".join(labels_above_threshold)
    # Preprocess image
    transform = transforms.Compose([
        transforms.Resize((224, 224)),
        transforms.ToTensor()
    ])
    image_tensor = transform(image).unsqueeze(0)

    #Convert prompt into embedding
    text_embedding = encode_prompt(question)

    #Model prediction
    with torch.no_grad():
        output = model(image_tensor, text_embedding)
        similarity_score = torch.sigmoid(output).item()

    #LLM Response
    prompt = (
    f"You are a highly knowledgeable radiologist interpreting a confirmed X-ray image. "
    f"The x-ray shows proof of the following: {label_text}. The findings in this are most likely true. \n\n"
    f"The patient (user) has asked the following question:\n\n"
    f"'{question}'\n\n"
    "You must answer solely based on the conditions listed above. "
    "Do not ask for additional tests, disclaim responsibility, or refer to human doctors. "
    "Assume the findings are accurate and confirmed. "
    "Do not suggest hypothetical signs or general radiology patterns. "
    "Do not say 'if I had the image' — you have the relevant information.\n\n"
    "Provide a structured, clear, and professional radiological interpretation in response to the question. "
    "Your tone must be confident, clinical, and based on the provided findings only."
)


    api_key = "C7Xwjta0ERFDWptomYooHrTOqMDWlS51"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"model": "mistral-small", "messages": [{"role": "user", "content": prompt}]}
    
    response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def load_medgemma():
    hf_token = "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"
    tokenizer = AutoTokenizer.from_pretrained("google/medgemma-4b-it", token=hf_token)
    model = AutoModel.from_pretrained(
        "google/medgemma-4b-it",
        token=hf_token,
        device_map="auto",
        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
    )
    return tokenizer, model


def generate_auto_report(labels: list[str]) -> BytesIO:
    label_text = ", ".join(labels)
    prompt = (
        f"Out of these medical conditions or radiological patterns:\n\n {class_names}\n\n"
        f"the X-ray image only shows signs of: {label_text}. \n"
        "Please write a short, professional radiology report which explains any medical conditions or radiological patterns that were identified."
        "If there are no abnormalities found, the x-ray shows that the chest is normal"

        "Do not include any sections about clinical history, don't repeat the title, "
        "and do not sign the report or add any closing statements like 'Sincerely'."
        "Exclude any sections or phrases related to contact, disclaimers, clinical history, greetings, or closings. "
        "Do not include phrases like 'please contact us', 'for further information', 'sincerely', or similar. "
        "The report should include only technical observations and a clinical impression."

    )

    api_key = "C7Xwjta0ERFDWptomYooHrTOqMDWlS51"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {"model": "mistral-small", "messages": [{"role": "user", "content": prompt}]}
    
    response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=data)
    response.raise_for_status()
    report_text = response.json()["choices"][0]["message"]["content"]

    # Create a Word document
    doc = Document()
    doc.add_heading("X-Ray Interpretation", level=1)
    doc.add_paragraph(report_text)

    # Save to in-memory file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer




def encode_prompt(prompt: str):
    tokens = tokenizer(prompt, return_tensors="pt", truncation=True, padding=True)
    with torch.no_grad():
        outputs = text_encoder(**tokens)
        embedding = outputs.last_hidden_state.mean(dim=1) 
        embedding_256 = embedding[:, :256]  
    return embedding_256


    
def get_multilabel_predictions(image: Image.Image, model: nn.Module, threshold=0.5):
    transform = transforms.Compose([
        transforms.Resize((224, 224)),
        transforms.ToTensor()
    ])
    tensor = transform(image.convert("RGB")).unsqueeze(0)

    with torch.no_grad():
        logits = model.encoder(tensor)
        probs = torch.sigmoid(logits).squeeze(0)

    predictions = [(class_names[i], probs[i].item()) for i in range(len(class_names))]
    return predictions



st.sidebar.button("Back to Main Menu")
st.title("X-ray Analyzer")

uploaded_file = st.file_uploader("Upload your X-ray image here:", type=["png", "jpg", "jpeg"])
if uploaded_file:
    image = Image.open(uploaded_file)
    if image.mode == "I;16":
        image = image.point(lambda i: i * (1.0 / 256)).convert("L")
    image = image.convert("L")
    image = Image.merge("RGB", (image, image, image))

    col1, col2 = st.columns(2)
    with col1:
        x_pixels = st.number_input("Width", min_value=1, step=1, value=224)
    with col2:
        y_pixels = st.number_input("Height", min_value=1, step=1, value=224)

    image = image.resize((x_pixels, y_pixels))
    st.session_state.image = image
    st.image(st.session_state.image, caption="Preview of Uploaded X-ray", width=300)


    #Display predicted conditions
    if "last_uploaded_filename" not in st.session_state or uploaded_file.name != st.session_state.last_uploaded_filename:
        st.markdown("### 🔍 Predicted Conditions:")
        st.session_state.predictions = get_multilabel_predictions(image, model)
        st.session_state.last_uploaded_filename = uploaded_file.name
    


    for label, score in st.session_state.predictions:
        emoji = "✅" if score >= 0.7 else "❌"
        st.write(f"- **{label}**: {score:.2f} {emoji}")



    if "auto_doc" not in st.session_state:
        st.session_state.auto_doc = None
    if "report_summary" not in st.session_state:
        st.session_state.report_summary = None

    if st.button("Generate Doctor's Report"):
        with st.spinner("Analyzing and generating doctor's report..."):
            labels_above_threshold = [label for label, score in st.session_state.predictions if score >= 0.7] #THRESHOLD DETERMINED HERE!!!!!!!!!!!!!!!!!
            auto_doc = generate_auto_report(labels_above_threshold)

            st.session_state.auto_doc = auto_doc
            st.session_state.report_summary = doctors_report_summary(auto_doc)
            st.success("Doctor's report is ready!")

    if st.session_state.auto_doc:
        buffer = BytesIO()
        buffer.seek(0)

        st.download_button(
            label="Download Doctor's Report",
            data=st.session_state.auto_doc,
            file_name="radiology_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )



        st.text_area("Summary of Doctor's Report:", value=st.session_state.report_summary, disabled=True)


    disable_text_input = False
    user_input_message = "Ask a follow-up question here"
else:
    disable_text_input = True
    user_input_message = "Upload image first"

user_input = st.text_area("Ask Questions Here", placeholder=user_input_message, value=st.session_state.selected_question or "", disabled=disable_text_input)

if st.button("Submit") and user_input and uploaded_file:
    with torch.no_grad():
        explanation = get_model_response(image, user_input, model)
        st.session_state.ai_response = explanation

    

    st.session_state.ai_response = explanation

    

st.text_area("AI Response", value=st.session_state.ai_response or "", height=150, disabled=True)

